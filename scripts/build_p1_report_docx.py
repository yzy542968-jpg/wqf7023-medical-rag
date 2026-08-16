from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "P1_FORMAL_REPORT.md"
OUTPUT = ROOT / "deliverables" / "p1-formal-report.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x2B, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)
TABLE_FILL = "F4F6F9"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_fixed_width(table, column_count: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total_width = 9360
    col_width = total_width // max(column_count, 1)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell in row.cells:
            _set_cell_width(cell, col_width)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.text = "WQF7023 MAI Research Project Proposal"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = MUTED


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)


def _add_paragraph_with_inline_code(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(8 if style is None else 4)
    chunks = re.split(r"(`[^`]+`)", text)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk.strip("`"))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(chunk)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_fixed_width(table, column_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(column_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            if r_idx == 0:
                run.font.bold = True
                _set_cell_shading(cell, TABLE_FILL)
    document.add_paragraph()


def build_docx() -> None:
    document = Document()
    _configure_document(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []
    in_code = False
    metadata_mode = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _add_table(document, _parse_table(table_buffer))
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            if line.strip():
                paragraph = document.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue

        if line.startswith("# "):
            _add_title(document, line[2:].strip())
            metadata_mode = True
            continue
        if line.startswith("## "):
            metadata_mode = False
            heading = line[3:].strip()
            document.add_heading(heading, level=1)
            continue
        if line.startswith("### "):
            metadata_mode = False
            document.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            metadata_mode = False
            document.add_heading(line[5:].strip(), level=3)
            continue
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            _add_paragraph_with_inline_code(document, text, style="List Number")
            continue
        if line.startswith("- "):
            _add_paragraph_with_inline_code(document, line[2:].strip(), style="List Bullet")
            continue
        if metadata_mode and ":" in line:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            label, value = line.split(":", 1)
            label_run = paragraph.add_run(label + ":")
            label_run.bold = True
            paragraph.add_run(value)
            continue

        _add_paragraph_with_inline_code(document, line)

    flush_table()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
