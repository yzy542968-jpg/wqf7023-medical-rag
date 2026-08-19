from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

import build_v5_manuscript_docx as manuscript


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/P2_V5_INTEGRATED_MANUSCRIPT.md"
OUTPUT = ROOT / "docs/P2_V5_UM_TEMPLATE_MANUSCRIPT.docx"
TEMPLATE_PATTERN = "**/P1_ZHANG_YUE_22097191_Final_01.docx"
TABLE_CAPTIONS = (
    ("Table 4.1. Retrieval results under four principal input conditions", "_Table41"),
    ("Table 4.2. End-to-end QA comparison", "_Table42"),
    ("Table 4.3. Runtime and computational cost", "_Table43"),
)
ABBREVIATIONS = (
    ("AI", "Artificial Intelligence"),
    ("BioViL-T", "BioViL with Transformer-based text encoding"),
    ("BM25", "Best Matching 25 probabilistic retrieval model"),
    ("CXR", "Chest X-ray"),
    ("EM", "Exact Match"),
    ("GPU", "Graphics Processing Unit"),
    ("Hit@k", "Proportion of queries retrieving the target within the top k ranks"),
    ("IU X-Ray/OpenI", "Indiana University Chest X-Ray/OpenI collection"),
    ("LLM", "Large Language Model"),
    ("MRR", "Mean Reciprocal Rank"),
    ("QA", "Question Answering"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("VLM", "Vision-Language Model"),
    ("VRAM", "Video Random Access Memory"),
)

ENGLISH_ABSTRACT = (
    "Retrieval-augmented generation can provide external evidence for medical question answering, but "
    "a fluent answer may still be grounded in the wrong patient's report. This study develops and "
    "critically evaluates a multimodal retrieval-augmented question-answering workflow over paired chest "
    "X-ray images and reports using de-identified IU X-Ray/OpenI cases. The frozen V5 experiment used "
    "240 fresh cases divided into 120 development and 120 confirmation cases, with 360 confirmation "
    "questions. BM25 provided text retrieval, BioViL-T supplied 128-dimensional paired image-report "
    "representations for reranking, and 100 fixed-point-free shuffled-image permutations provided an "
    "alignment control. The top-ranked report was passed through a fixed local Qwen2.5-1.5B-Instruct "
    "generator and an automated Medical-NLI evidence checker. Indication text was the strongest retrieval "
    "signal: MRR increased from 0.0277 for question-only BM25 to 0.6590 for indication-plus-question "
    "BM25. Adding the correctly aligned image increased MRR to 0.6971 and extractive proxy Token-F1 from "
    "0.6602 to 0.7245. The MRR difference was 0.0381 with case-bootstrap 95% CI [0.0159, 0.0614] and "
    "paired-randomization p=0.0012. No shuffled-image run reached the correctly aligned MRR or proxy "
    "Token-F1; the plus-one Monte Carlo value was 0.0099 for both metrics. Under the fixed downstream "
    "pipeline, multimodal retrieval improved final Token-F1 by 0.0302, CI [0.0101, 0.0511], p=0.0032, "
    "but automated support rate decreased by 0.0340. A frozen 24-question researcher-reviewed qualitative "
    "analysis showed that target-rank improvement did not always produce Top-1 success, report-level "
    "faithfulness did not guarantee target-case alignment, correct retrieval did not guarantee a "
    "reference-consistent answer, and automated verification sometimes appeared to remove report-supported "
    "content. The results support an alignment-specific image contribution to paired-report retrieval, "
    "but they do not establish autonomous image diagnosis, clinical correctness, external validation, or "
    "deployment safety."
)

MALAY_ABSTRACT = (
    "Penjanaan berasaskan dapatan semula boleh menyediakan bukti luaran untuk penjawaban soalan perubatan, "
    "namun jawapan yang lancar masih boleh diasaskan pada laporan pesakit yang salah. Kajian ini membangunkan "
    "dan menilai secara kritikal aliran kerja penjawaban soalan berbilang modal berasaskan dapatan semula "
    "menggunakan pasangan imej X-ray dada dan laporan IU X-Ray/OpenI yang dinyahpengenalan. Eksperimen V5 "
    "beku menggunakan 240 kes baharu yang dibahagikan kepada 120 kes pembangunan dan 120 kes pengesahan, "
    "dengan 360 soalan pengesahan. BM25 digunakan untuk dapatan semula teks, BioViL-T menghasilkan perwakilan "
    "imej-laporan 128 dimensi untuk penyusunan semula, dan 100 permutasi imej rawak tanpa titik tetap digunakan "
    "sebagai kawalan penjajaran. Laporan kedudukan teratas diproses oleh penjana Qwen2.5-1.5B-Instruct tempatan "
    "yang tetap serta pemeriksa bukti Medical-NLI automatik. Teks petunjuk merupakan isyarat dapatan semula "
    "terkuat: MRR meningkat daripada 0.0277 bagi BM25 soalan sahaja kepada 0.6590 bagi BM25 petunjuk bersama "
    "soalan. Penambahan imej yang dipadankan dengan betul meningkatkan MRR kepada 0.6971 dan Token-F1 proksi "
    "ekstraktif daripada 0.6602 kepada 0.7245. Perbezaan MRR ialah 0.0381 dengan 95% CI bootstrap kes "
    "[0.0159, 0.0614] dan p=0.0012 bagi rawakan berpasangan. Tiada larian imej rawak mencapai MRR atau Token-F1 "
    "proksi bagi imej yang dipadankan dengan betul; nilai Monte Carlo tambah-satu ialah 0.0099 untuk kedua-dua "
    "metrik. Dalam saluran hiliran yang tetap, dapatan semula berbilang modal meningkatkan Token-F1 akhir "
    "sebanyak 0.0302, CI [0.0101, 0.0511], p=0.0032, tetapi kadar sokongan automatik menurun sebanyak 0.0340. "
    "Analisis kualitatif beku terhadap 24 soalan yang disemak oleh penyelidik menunjukkan bahawa peningkatan "
    "kedudukan sasaran tidak semestinya menghasilkan kejayaan Top-1, kesetiaan pada laporan tidak menjamin "
    "penjajaran dengan kes sasaran, dapatan semula yang betul tidak menjamin jawapan yang konsisten dengan "
    "rujukan, dan pengesahan automatik kadangkala kelihatan membuang kandungan yang disokong laporan. Hasil ini "
    "menyokong sumbangan imej yang khusus kepada penjajaran bagi dapatan semula laporan berpasangan, tetapi tidak "
    "membuktikan diagnosis imej autonomi, ketepatan klinikal, pengesahan luaran, atau keselamatan penggunaan."
)


def find_template() -> Path:
    base = Path.home() / "Desktop" / "2025-2026-s2" / "P1"
    matches = sorted(base.glob(TEMPLATE_PATTERN))
    if not matches:
        raise FileNotFoundError(f"Could not locate template matching {TEMPLATE_PATTERN}")
    return matches[0]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def insert_paragraph_after(paragraph, text: str = "") -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._element.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if text:
        inserted.add_run(text)
    return inserted


def add_bookmark(paragraph: Paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def remove_original_body(doc: Document) -> None:
    chapter = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "CHAPTER 1: INTRODUCTION"
    )
    body = doc._body._element
    start = list(body).index(chapter._element)
    for child in list(body)[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def rebuild_front_matter(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    replace_paragraph_text(paragraphs[73], ENGLISH_ABSTRACT)
    replace_paragraph_text(
        paragraphs[74],
        "Keywords: retrieval-augmented generation, multimodal retrieval, radiology, medical question answering, evidence grounding",
    )
    replace_paragraph_text(paragraphs[78], MALAY_ABSTRACT)
    replace_paragraph_text(
        paragraphs[79],
        "Kata kunci: penjanaan berasaskan dapatan semula, dapatan semula berbilang modal, radiologi, penjawaban soalan perubatan, pembumian bukti",
    )
    replace_paragraph_text(
        paragraphs[82],
        "I would like to express my sincere appreciation to my supervisor, Dr. Uzair Ishtiaq, for his guidance and support throughout this research project. I am also grateful to the Faculty of Computer Science and Information Technology, Universiti Malaya, for providing the academic environment and resources that supported this work.",
    )

    toc_heading = next(p for p in doc.paragraphs if p.text.strip() == "TABLE OF CONTENTS")
    symbols_heading = next(p for p in doc.paragraphs if p.text.strip() == "LIST OF SYMBOLS AND ABBREVIATIONS")
    body = doc._body._element
    toc_idx = list(body).index(toc_heading._element)
    symbols_idx = list(body).index(symbols_heading._element)
    for child in list(body)[toc_idx + 1 : symbols_idx]:
        body.remove(child)

    toc_paragraph = OxmlElement("w:p")
    toc_heading._element.addnext(toc_paragraph)
    toc = next(p for p in doc.paragraphs if p._element is toc_paragraph)
    manuscript.add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Update this table of contents in Microsoft Word.")

    page_break = insert_paragraph_after(toc)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    list_heading = insert_paragraph_after(page_break, "LIST OF TABLES")
    list_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    list_heading.runs[0].bold = True
    cursor = list_heading
    for caption, bookmark in TABLE_CAPTIONS:
        cursor = insert_paragraph_after(cursor)
        cursor.add_run(caption)
        cursor.add_run("\t")
        manuscript.add_field(cursor, f"PAGEREF {bookmark} \\h")

    abbreviation_table = doc.tables[0]
    for row in list(abbreviation_table.rows)[1:]:
        abbreviation_table._tbl.remove(row._tr)
    for abbreviation, meaning in ABBREVIATIONS:
        cells = abbreviation_table.add_row().cells
        cells[0].text = abbreviation
        cells[1].text = meaning


def configure_um_body(doc: Document) -> None:
    section = doc.sections[-1]
    content_width = section.page_width - section.left_margin - section.right_margin
    manuscript.CONTENT_DXA = int(content_width / 635)
    manuscript.TABLE_INDENT_DXA = 120
    for table in doc.tables:
        column_count = len(table.columns)
        widths = [manuscript.CONTENT_DXA // column_count] * column_count
        widths[-1] += manuscript.CONTENT_DXA - sum(widths)
        manuscript.set_table_geometry(table, widths)
    if "Code Block" not in doc.styles:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)
    manuscript.set_update_fields(doc)


def main() -> None:
    template = find_template()
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document(template)
    remove_original_body(doc)
    rebuild_front_matter(doc)
    configure_um_body(doc)
    manuscript.add_body(doc, markdown, start_heading="# Chapter 1: Introduction")

    bookmark_id = 700
    caption_map = {caption: bookmark for caption, bookmark in TABLE_CAPTIONS}
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.startswith("Chapter "):
            replace_paragraph_text(paragraph, paragraph.text.upper())
        bookmark = caption_map.get(paragraph.text.strip())
        if bookmark:
            add_bookmark(paragraph, bookmark, bookmark_id)
            bookmark_id += 1

    props = doc.core_properties
    props.title = "Retrieval-Augmented Medical Question Answering over Paired Radiology Images and Reports"
    props.author = "Zhang Yue"
    props.subject = "Master of Artificial Intelligence Research Project"
    props.comments = "UM-template V5-integrated manuscript; technical and qualitative results frozen."
    doc.save(OUTPUT)
    print(f"Template: {template}")
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
