from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_FILES = [
    ROOT / "docs" / "P2_V6_FRONT_MATTER_CH1_2.md",
    ROOT / "docs" / "P2_V6_INTEGRATED_CHAPTERS_3_5.md",
]
DEFAULT_OUTPUT = ROOT / "docs" / "P2_V6_TECHNICAL_INTEGRATION_DRAFT.docx"


def set_run_font(run, name: str = "Calibri", size: float = 11, bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=9)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("V6 integrated technical draft | ")
    set_run_font(run, size=9)
    run.font.color.rgb = RGBColor.from_string("666666")
    add_page_number(footer)


def inline_runs(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(9, size - 1))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows[0], rows[1:]


def add_table(doc: Document, lines: list[str]) -> None:
    headers, body = parse_table(lines)
    rows = [headers, *body]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    if len(headers) == 5:
        widths = [2650, 1450, 1450, 1450, 2360]
    elif len(headers) == 4:
        widths = [2350, 1650, 2400, 2960]
    elif len(headers) == 3:
        widths = [2400, 2300, 4660]
    else:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            inline_runs(paragraph, value, size=9)
            if row_index == 0:
                set_cell_shading(cell, "F4F6F9")
                for run in paragraph.runs:
                    run.bold = True
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown(doc: Document, text: str) -> None:
    lines = text.replace("\r\n", "\n").split("\n")
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.right_indent = Inches(0.2)
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.0
                for code_line in code_lines:
                    run = paragraph.add_run(code_line + "\n")
                    set_run_font(run, name="Consolas", size=9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Heading 1")
            inline_runs(paragraph, line[2:].strip(), size=16)
            index += 1
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 2")
            inline_runs(paragraph, line[3:].strip(), size=13)
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            inline_runs(paragraph, line[4:].strip(), size=12)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            inline_runs(paragraph, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            inline_runs(paragraph, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=11)
            run.italic = True
            index += 1
            continue
        paragraph = doc.add_paragraph()
        inline_runs(paragraph, line.strip())
        index += 1


def main() -> None:
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT
    document = Document()
    configure_document(document)
    for source_index, source in enumerate(SOURCE_FILES):
        add_markdown(document, source.read_text(encoding="utf-8"))
        if source_index == 0:
            document.add_page_break()
    document.core_properties.title = "V6 Integrated Technical Draft"
    document.core_properties.subject = "V6 model-modernized confirmation study"
    document.core_properties.author = "ZHANG YUE"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
