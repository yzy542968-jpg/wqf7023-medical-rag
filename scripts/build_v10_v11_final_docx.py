from __future__ import annotations

import re
from math import sqrt
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "P2_V10_V11_FINAL_MANUSCRIPT.md"
OUTPUT = ROOT / "deliverables" / "22097191_ZHANG_YUE_Final_Research_Project.docx"

BLUE = RGBColor(31, 78, 121)
DARK_BLUE = RGBColor(11, 37, 69)
MUTED = RGBColor(91, 103, 112)
LIGHT = "F4F6F9"
CONTENT_WIDTH_DXA = 9026  # A4 width minus two 1-inch margins.


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def allocate_table_widths(rows: list[list[str]]) -> list[int]:
    """Allocate A4-safe column widths from bounded content-length weights."""
    cols = len(rows[0])
    weights: list[float] = []
    for column in range(cols):
        longest = max(len(row[column]) for row in rows)
        weights.append(max(2.5, min(7.5, sqrt(max(longest, 1)))))

    raw = [CONTENT_WIDTH_DXA * weight / sum(weights) for weight in weights]
    widths = [max(760, int(value)) for value in raw]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[weights.index(max(weights))] += difference
    return widths


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def apply_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            apply_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            apply_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color=DARK_BLUE)
        elif token.startswith("**"):
            apply_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            apply_font(paragraph.add_run(token[1:-1]), italic=True)
        cursor = match.end()
    if cursor < len(text):
        apply_font(paragraph.add_run(text[cursor:]))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Code Block" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = DARK_BLUE
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "false")


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("RETRIEVAL-AUGMENTED MEDICAL QUESTION ANSWERING\nOVER PAIRED RADIOLOGY IMAGES AND REPORTS")
    apply_font(r, size=20, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("Final V10 Primary Study with V11 Development Extension")
    apply_font(r, size=14, italic=True, color=MUTED)

    for label, value in (
        ("Student", "ZHANG YUE"),
        ("Student ID", "22097191"),
        ("Programme", "Master of Artificial Intelligence"),
        ("Institution", "University of Malaya"),
        ("Date", "August 2026"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        apply_font(p.add_run(f"{label}: "), size=11, bold=True)
        apply_font(p.add_run(value), size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("Retrospective research prototype. Not clinically validated.")
    apply_font(r, size=10, italic=True, color=MUTED)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    add_inline(p, "Table of Contents")
    for item, page_number in (
        ("Abstract", 3),
        ("Chapter 1: Introduction", 5),
        ("Chapter 2: Literature Review", 10),
        ("Chapter 3: Methodology", 21),
        ("Chapter 4: Results and Analysis", 27),
        ("Chapter 5: Discussion", 32),
        ("Chapter 6: Conclusion", 36),
        ("References", 37),
        ("Appendices", 39),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.05),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS,
        )
        add_inline(p, f"{item}\t{page_number}")
    doc.add_page_break()


def add_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    apply_font(p.add_run("V10/V11 FINAL RESEARCH MANUSCRIPT"), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    apply_font(p.add_run("22097191 | "), size=8.5, color=MUTED)
    add_field(p, "PAGE")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    cleaned = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(cleaned), cols=cols)
    table.style = "Table Grid"
    widths = allocate_table_widths(cleaned)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, values in enumerate(cleaned):
        for cidx, value in enumerate(values):
            cell = table.cell(ridx, cidx)
            set_cell_margins(cell)
            if ridx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value)
            for run in p.runs:
                apply_font(run, size=9, bold=(ridx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("suff", "tab")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    add_inline(p, text)


def parse_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 1  # skip document title, already rendered on the cover
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    active_num_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline(p, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            active_num_id = None
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                add_inline(p, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        if not stripped:
            flush_paragraph()
            active_num_id = None
            idx += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            active_num_id = None
            title = stripped[2:]
            p = doc.add_paragraph(style="Heading 1")
            if title.startswith("Chapter ") or title in {"References", "Appendices"}:
                p.paragraph_format.page_break_before = True
            add_inline(p, title)
            idx += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[3:])
            idx += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[4:])
            idx += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            active_num_id = None
            table_rows: list[list[str]] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                row = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
                    table_rows.append(row)
                idx += 1
            add_table(doc, table_rows)
            continue
        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            flush_paragraph()
            if active_num_id is None:
                active_num_id = create_numbering_instance(doc)
            add_numbered_paragraph(doc, number_match.group(2), active_num_id)
            idx += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            idx += 1
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            active_num_id = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            add_inline(p, stripped[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = DARK_BLUE
            idx += 1
            continue
        active_num_id = None
        paragraph_buffer.append(stripped)
        idx += 1
    flush_paragraph()


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_running_furniture(doc)
    add_cover(doc)
    parse_markdown(doc, markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
