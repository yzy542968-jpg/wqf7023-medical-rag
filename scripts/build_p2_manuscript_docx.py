from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "P2_FINAL_MANUSCRIPT.md"
OUTPUT = ROOT / "deliverables" / "22097191_ZHANG_YUE_P2_Research_Project.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360

NAVY = RGBColor(0x17, 0x31, 0x4D)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x2B, 0x33)
MUTED = RGBColor(0x65, 0x6D, 0x76)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = "EAF2F8"
TABLE_HEADER = "DCE6F1"
TABLE_ALT = "F7F9FB"
CODE_FILL = "F3F5F7"


def _set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def _set_style_font(style, name: str, size: float, color: RGBColor | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            _set_cell_margins(cell)


def _set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend((begin, instr, separate))
    if placeholder:
        value_run = paragraph.add_run(placeholder)
        _set_run_font(value_run, "Calibri", 9)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 6),
        "Heading 3": (11.5, DARK_BLUE, 10, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        _set_style_font(style, "Calibri", size, color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        _set_style_font(style, "Calibri", 11, INK)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    document.settings.element.append(OxmlElement("w:updateFields"))
    document.settings.element[-1].set(qn("w:val"), "true")


def _configure_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("WQF7023 MAI Research Project")
    _set_run_font(run, "Calibri", 8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("ZHANG YUE  |  22097191  |  ")
    _set_run_font(run, "Calibri", 8.5)
    run.font.color.rgb = MUTED
    _add_field(footer, " PAGE ", "1")


def _add_cover(document: Document, title: str, metadata: list[tuple[str, str]]) -> None:
    section = document.sections[0]
    _configure_page(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)

    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(18)
    run = eyebrow.add_run("MASTER OF ARTIFICIAL INTELLIGENCE")
    _set_run_font(run, "Calibri", 10)
    run.font.bold = True
    run.font.color.rgb = BLUE

    title_p = document.add_paragraph()
    title_p.paragraph_format.space_after = Pt(18)
    title_p.paragraph_format.line_spacing = 1.05
    title_p.paragraph_format.keep_with_next = True
    title_run = title_p.add_run(title)
    _set_run_font(title_run, "Calibri", 24)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(34)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    p_pr.append(borders)

    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.0
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, "Calibri", 10.5)
        label_run.font.bold = True
        label_run.font.color.rgb = MUTED
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, "Calibri", 10.5)
        value_run.font.color.rgb = INK

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(32)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("FINAL RESEARCH MANUSCRIPT - AUTOMATED RESULTS FROZEN")
    _set_run_font(run, "Calibri", 9)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE


def _add_toc(document: Document) -> None:
    heading = document.add_heading("Table of Contents", level=1)
    heading.paragraph_format.space_before = Pt(0)
    entries = [
        ("Abstract", 2),
        ("Chapter 1: Introduction", 3),
        ("Chapter 2: Literature Review", 6),
        ("Chapter 3: Methodology", 9),
        ("Chapter 4: Results and Analysis", 13),
        ("Chapter 5: Discussion and Conclusion", 17),
        ("References", 20),
        ("Appendices", 21),
    ]
    for label, page in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.45), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = paragraph.add_run(f"{label}\t{page}")
        _set_run_font(run, "Calibri", 10.5)
        run.font.color.rgb = INK


def _new_decimal_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "440")
    indentation.set(qn("w:hanging"), "260")
    p_pr.append(indentation)
    for node in (start, number_format, level_text, justification, p_pr):
        level.append(node)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_inline(paragraph, text: str) -> None:
    for chunk in INLINE_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            _set_run_font(run, "Consolas", 9.5)
        elif chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        else:
            paragraph.add_run(chunk)


def _add_body_paragraph(document: Document, text: str, style: str | None = None, *, reference: bool = False):
    paragraph = document.add_paragraph(style=style)
    if style is None:
        paragraph.paragraph_format.first_line_indent = Inches(0.25)
    if reference:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(6)
    _add_inline(paragraph, text)
    return paragraph


def _add_blockquote(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_widths(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, LIGHT_BLUE)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "bottom", "end"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "18")
    start.set(qn("w:color"), "2E74B5")
    borders.append(start)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_inline(paragraph, text)
    for run in paragraph.runs:
        run.italic = True
        run.font.color.rgb = NAVY
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def _column_widths(rows: list[list[str]], column_count: int) -> list[int]:
    scores: list[int] = []
    for index in range(column_count):
        lengths = [len(row[index]) if index < len(row) else 0 for row in rows]
        scores.append(max(8, min(max(lengths, default=8), 42)))
    minimum = 1500 if column_count <= 4 else 1000
    remaining = CONTENT_WIDTH_DXA - minimum * column_count
    if remaining <= 0:
        return [CONTENT_WIDTH_DXA // column_count] * column_count
    score_total = sum(scores)
    widths = [minimum + round(remaining * score / score_total) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = _column_widths(rows, column_count)
    _set_table_widths(table, widths)

    for row_index, row in enumerate(rows):
        if row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            table.rows[0]._tr.get_or_add_trPr().append(tbl_header)
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                _set_cell_shading(cell, TABLE_HEADER)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, TABLE_ALT)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            text = row[column_index] if column_index < len(row) else ""
            _add_inline(paragraph, text)
            for run in paragraph.runs:
                _set_run_font(run, "Calibri", 8.7)
                if row_index == 0:
                    run.font.bold = True
                    run.font.color.rgb = NAVY
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def _add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        _set_run_font(run, "Consolas", 8.5)
        run.font.color.rgb = NAVY
        if index < len(lines) - 1:
            run.add_break()


def _extract_title_metadata(lines: list[str]) -> tuple[str, list[tuple[str, str]]]:
    title = lines[0][2:].strip()
    metadata: list[tuple[str, str]] = []
    for line in lines[1:]:
        if line.startswith("## "):
            break
        clean = line.strip().rstrip()
        if not clean or ":" not in clean:
            continue
        clean = clean.replace("**", "")
        label, value = clean.split(":", 1)
        metadata.append((label.strip(), value.strip()))
    return title, metadata


def build_docx() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title, metadata = _extract_title_metadata(lines)

    document = Document()
    _configure_styles(document)
    _add_cover(document, title, metadata)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_page(body_section)
    _set_page_number_start(body_section, 1)
    _configure_header_footer(body_section)
    _add_toc(document)
    document.add_page_break()

    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    content_started = False
    in_references = False
    active_num_id: int | None = None

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _add_table(document, _parse_table(table_buffer))
            table_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        _add_code_block(document, code_buffer)
        code_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if not content_started:
            if line == "## Abstract":
                content_started = True
            else:
                continue

        if line.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue

        if line.startswith("# "):
            active_num_id = None
            heading_text = line[2:].strip()
            heading = document.add_heading(heading_text, level=1)
            if heading_text.startswith("Chapter ") or heading_text in {"References", "Appendices"}:
                heading.paragraph_format.page_break_before = True
            in_references = heading_text == "References"
            continue
        if line.startswith("## "):
            active_num_id = None
            heading_text = line[3:].strip()
            in_references = False
            document.add_heading(heading_text, level=1 if heading_text == "Abstract" else 2)
            continue
        if line.startswith("### "):
            active_num_id = None
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("> "):
            active_num_id = None
            _add_blockquote(document, line[2:].strip())
            continue
        if re.match(r"^\d+\.\s+", line):
            if active_num_id is None:
                active_num_id = _new_decimal_numbering(document)
            paragraph = _add_body_paragraph(
                document, re.sub(r"^\d+\.\s+", "", line), style="List Number"
            )
            _apply_numbering(paragraph, active_num_id)
            continue
        active_num_id = None
        if line.startswith("- "):
            _add_body_paragraph(document, line[2:].strip(), style="List Bullet")
            continue

        _add_body_paragraph(document, line, reference=in_references)

    flush_table()
    if in_code:
        flush_code()

    core = document.core_properties
    core.title = title
    core.subject = "WQF7023 Artificial Intelligence Research Project"
    core.author = "ZHANG YUE"
    core.keywords = "RAG, radiology, evidence checking, patient scope, medical question answering"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_docx())
