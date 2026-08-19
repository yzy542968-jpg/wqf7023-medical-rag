from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/P2_V5_INTEGRATED_MANUSCRIPT.md"
OUTPUT = ROOT / "docs/P2_V5_INTEGRATED_MANUSCRIPT.docx"

FONT = "Calibri"
BODY_SIZE = 11
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
MUTED = "666666"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None, color: str | None = None, name: str = FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_field(paragraph, instruction: str, display: str = "1") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, separate, text, end):
        run._r.append(node)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(BODY_SIZE)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    quote = styles["Quote"]
    quote.font.name = FONT
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.right_indent = Inches(0.25)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.right_indent = Inches(0.15)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    if "TOC Heading Custom" not in styles:
        toc_heading = styles.add_style("TOC Heading Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc_heading = styles["TOC Heading Custom"]
    toc_heading.font.name = FONT
    toc_heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    toc_heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    toc_heading.font.size = Pt(18)
    toc_heading.font.bold = True
    toc_heading.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after = Pt(14)
    toc_heading.paragraph_format.keep_with_next = True


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str) -> None:
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token)
            set_run_font(run)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0]] + rows[2:]
    cols = len(rows[0])
    if any(len(row) != cols for row in rows):
        raise ValueError("Inconsistent Markdown table")
    if cols == 6:
        widths = [3000, 1050, 1050, 1050, 1050, 2160]
    elif cols == 5:
        widths = [3200, 1540, 1540, 1540, 1540]
    elif cols == 4:
        widths = [3500, 1953, 1953, 1954]
    elif cols == 3:
        widths = [3000, 3180, 3180]
    else:
        widths = [CONTENT_DXA // cols] * cols
        widths[-1] += CONTENT_DXA - sum(widths)

    if sum(widths) != CONTENT_DXA:
        source_total = sum(widths)
        widths = [round(width * CONTENT_DXA / source_total) for width in widths]
        widths[-1] += CONTENT_DXA - sum(widths)

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value)
            for run in p.runs:
                set_run_font(run, size=9, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
    repeat_table_header(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RETRIEVAL-AUGMENTED MEDICAL QUESTION ANSWERING\nOVER PAIRED RADIOLOGY IMAGES AND REPORTS")
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(46)
    run = sub.add_run("WQF7023 Artificial Intelligence Research Project")
    set_run_font(run, size=13, bold=True, color=MUTED)

    for text, bold in (
        ("ZHANG YUE", True),
        ("Matric No. 22097191", False),
        ("Master of Artificial Intelligence", False),
        ("Supervisor: Dr. Uzair Ishtiaq", False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=bold)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run("19 August 2026")
    set_run_font(run, size=11, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph(style="TOC Heading Custom")
    p.add_run("Table of Contents")
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Right-click and update field to populate the table of contents.")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("The table of contents updates automatically when opened in Microsoft Word.")
    set_run_font(run, size=9, italic=True, color=MUTED)
    doc.add_page_break()


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("WQF7023 | Zhang Yue | V5-Integrated Manuscript")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_body(doc: Document, markdown: str, start_heading: str = "## Abstract") -> None:
    body = markdown.split(start_heading, maxsplit=1)[1]
    lines = [start_heading] + body.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(style="Code Block")
            add_inline(p, line or " ")
            for run in p.runs:
                set_run_font(run, name="Consolas", size=9)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|---"):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("# Chapter") or line in ("# References", "# Appendices"):
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.page_break_before = True
            add_inline(p, line[2:])
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            if line == "## Abstract":
                p.paragraph_format.page_break_before = False
            add_inline(p, line[3:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[4:])
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            add_inline(p, line[2:])
        elif re.match(r"^- ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            style = "Caption" if re.match(r"^Table \d+\.\d+", line) else None
            p = doc.add_paragraph(style=style)
            if line.startswith("Bae, ") or (doc.paragraphs and any(line.startswith(prefix) for prefix in ("Bannur,", "Boecking,", "Demner-", "Es,", "Jin,", "Lau,", "Lewis,", "Ngo,", "Pal,", "Qwen ", "Radford,", "Robertson,", "Romanov,", "Singhal,", "Soni,", "Xiong,"))):
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, line)
        i += 1


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)
    add_body(doc, markdown)
    set_update_fields(doc)

    props = doc.core_properties
    props.title = "Retrieval-Augmented Medical Question Answering over Paired Radiology Images and Reports"
    props.author = "Zhang Yue"
    props.subject = "WQF7023 Artificial Intelligence Research Project"
    props.keywords = "medical RAG, radiology, multimodal retrieval, BioViL-T"
    props.comments = "V5-integrated manuscript; technical and qualitative results frozen."

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
